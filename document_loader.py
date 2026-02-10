"""
Document Loader - Processes and chunks documents for RAG
"""
from typing import List, Dict, Any
import os
from pathlib import Path
from loguru import logger

# Document parsers
from pypdf import PdfReader
from docx import Document as DocxDocument
import pandas as pd

from config import settings


class DocumentLoader:
    """Loads and processes documents into chunks"""
    
    def __init__(self):
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """
        Load a document and extract text
        
        Args:
            file_path: Path to the document
        
        Returns:
            Dictionary with document content and metadata
        """
        file_extension = Path(file_path).suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return self._load_pdf(file_path)
            elif file_extension == '.docx':
                return self._load_docx(file_path)
            elif file_extension == '.txt':
                return self._load_txt(file_path)
            elif file_extension == '.csv':
                return self._load_csv(file_path)
            elif file_extension == '.md':
                return self._load_markdown(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
        
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {e}")
            raise
    
    def _load_pdf(self, file_path: str) -> Dict[str, Any]:
        """Load PDF document"""
        reader = PdfReader(file_path)
        text_parts = []
        
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                text_parts.append({
                    'text': text,
                    'page': page_num
                })
        
        return {
            'content': text_parts,
            'type': 'pdf',
            'total_pages': len(reader.pages)
        }
    
    def _load_docx(self, file_path: str) -> Dict[str, Any]:
        """Load DOCX document"""
        doc = DocxDocument(file_path)
        text_parts = []
        
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                text_parts.append({
                    'text': para.text,
                    'paragraph': i + 1
                })
        
        return {
            'content': text_parts,
            'type': 'docx',
            'total_paragraphs': len(doc.paragraphs)
        }
    
    def _load_txt(self, file_path: str) -> Dict[str, Any]:
        """Load TXT document"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        return {
            'content': [{'text': text, 'section': 1}],
            'type': 'txt'
        }
    
    def _load_csv(self, file_path: str) -> Dict[str, Any]:
        """Load CSV document"""
        df = pd.read_csv(file_path)
        
        # Convert to text representation
        text_parts = []
        text_parts.append({
            'text': f"CSV Headers: {', '.join(df.columns.tolist())}",
            'row': 0
        })
        
        for idx, row in df.iterrows():
            row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
            text_parts.append({
                'text': row_text,
                'row': idx + 1
            })
        
        return {
            'content': text_parts,
            'type': 'csv',
            'total_rows': len(df)
        }
    
    def _load_markdown(self, file_path: str) -> Dict[str, Any]:
        """Load Markdown document"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        return {
            'content': [{'text': text, 'section': 1}],
            'type': 'markdown'
        }
    
    def chunk_document(
        self,
        document: Dict[str, Any],
        filename: str
    ) -> List[Dict[str, Any]]:
        """
        Split document into chunks for vector storage
        
        Args:
            document: Document dictionary from load_document
            filename: Source filename
        
        Returns:
            List of chunks with metadata
        """
        chunks = []
        
        for part in document['content']:
            text = part['text']
            
            # Split into chunks with overlap
            text_chunks = self._split_text(text)
            
            for i, chunk_text in enumerate(text_chunks):
                chunk = {
                    'content': chunk_text,
                    'metadata': {
                        'source': filename,
                        'type': document['type'],
                        'chunk_index': i,
                        **{k: v for k, v in part.items() if k != 'text'}
                    }
                }
                chunks.append(chunk)
        
        logger.info(f"Created {len(chunks)} chunks from {filename}")
        return chunks
    
    def _split_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence ending
                for punct in ['. ', '! ', '? ', '\n\n']:
                    last_punct = text.rfind(punct, start, end)
                    if last_punct != -1:
                        end = last_punct + len(punct)
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - self.chunk_overlap
            if start < 0:
                start = 0
        
        return chunks
    
    def process_file(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Complete processing pipeline: load and chunk
        
        Args:
            file_path: Path to document
        
        Returns:
            List of chunks ready for vector store
        """
        filename = os.path.basename(file_path)
        logger.info(f"Processing document: {filename}")
        
        # Load document
        document = self.load_document(file_path)
        
        # Chunk document
        chunks = self.chunk_document(document, filename)
        
        return chunks


# Singleton instance
document_loader = DocumentLoader()
